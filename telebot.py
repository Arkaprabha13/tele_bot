from dotenv import load_dotenv
import os
from aiogram import Bot, Dispatcher, executor, types
import openai
import sys
from hugchat import hugchat
from hugchat.login import Login
import time
from docx import Document

class Reference:
    '''
    This class is used to store the reference to the bot and the dispatcher
    '''
    
    def __init__(self) -> None:
        self.response=""

load_dotenv()
TOKEN=os.getenv("TOKEN")
openai.api_key=os.getenv("OPEN_API_KEY")

reference= Reference()

model_name="gpt-3.5-turbo"

bot= Bot(token=TOKEN)
dispatcher=  Dispatcher(bot)

def clear_past_response():
    '''
    This function is used to clear the past response of the bot
    '''
    reference.response=""

def save_to_docx(question, answer, response_time):
    '''
    Save question-answer pair and response time to a Word document.
    '''
    doc = Document()
    doc.add_heading('Question-Answer Pairs', 0)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Question'
    hdr_cells[1].text = 'Answer'
    hdr_cells[2].text = 'Response Time (seconds)'
    row = table.add_row().cells
    row[0].text = question
    row[1].text = answer
    row[2].text = str(response_time)
    doc.save('question_answer_pairs.docx')

@dispatcher.message_handler(commands=['start'])
async def welcome(message: types.Message):
    await message.reply("Welcome to this bot!\nPowered by aiogram.")

@dispatcher.message_handler(commands=['clear'])
async def welcome(message: types.Message):
    clear_past_response()
    await message.reply("The past response has been cleared.")

@dispatcher.message_handler(commands=['help'])
async def welcome(message: types.Message):
    await message.reply("This bot is powered by OpenAI's GPT-3.5 model. You can ask any question and the bot will try to answer it.\n /start: Start the bot\n /help: Get help\n /clear: Clear the past response")

@dispatcher.message_handler()
async def all_time(message: types.Message):
    query=message.text
    print(f" >>> User: \t {query}")
    try:
        await message.answer("Please wait while I process your request...")
        start_time = time.time()
        sign = Login("arka13", "Arkaprabha13")
        cookies = sign.login()
        chatbot=hugchat.ChatBot(cookies=cookies.get_dict())
        id=chatbot.new_conversation()
        chatbot.change_conversation(id)
        print("Before chatbot.chat call")
        bot_response=chatbot.chat(query)
        print("After chatbot.chat call")
        print(f"\n <<< Bot: \t\n {bot_response}")
        end_time = time.time()
        await bot.send_message(chat_id=message.chat.id, text=f"Response Time: {end_time - start_time} seconds\n{bot_response}")
        save_to_docx(query, bot_response, end_time - start_time)
    except Exception as e:  
        print(e)
        await message.answer("Sorry, I am not able to answer that.")

if __name__ == '__main__':
    executor.start_polling(dispatcher, skip_updates=True)
